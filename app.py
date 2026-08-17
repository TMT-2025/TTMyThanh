import os
from flask import Flask, request, render_template, send_file, jsonify
import docx
import tempfile
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

app = Flask(__name__, template_folder='templates', static_folder='static')

WORKSPACE = os.path.dirname(os.path.abspath(__file__))
TEMPLATE_PATH = os.path.join(WORKSPACE, "BANG TUAN CUU 2.docx")
OUTPUT_PATH = os.path.join(tempfile.gettempdir(), "KetQuaTuanCuu.docx")

def set_cell_text_and_style(cell, text, font_name="Times New Roman", font_size=Pt(16), bold=False, align=WD_ALIGN_PARAGRAPH.CENTER):
    p = cell.paragraphs[0]
    p.text = ""
    p.alignment = align
    run = p.add_run(text)
    run.font.name = font_name
    run.font.size = font_size
    run.bold = bold

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/generate', methods=['POST'])
def generate():
    try:
        if request.is_json:
            data = request.json
        else:
            data = request.form
        hoTen = data.get('hoTen')
        phamDao = data.get('phamDao')
        gioMat = data.get('gioMat')
        ngayMatAl = data.get('ngayMatAl')
        ngayMatDl = data.get('ngayMatDl')
        
        tuoiMat_raw = data.get('tuoiMat')
        try:
            tuoiMat = int(tuoiMat_raw) if tuoiMat_raw else 0
        except ValueError:
            tuoiMat = 0
            
        namSinhAl = data.get('namSinhAl')
        tuoiTerm = data.get('tuoiTerm')
        
        tableData_raw = data.get('tableData')
        import json
        if isinstance(tableData_raw, str):
            try:
                tableData = json.loads(tableData_raw)
            except Exception:
                tableData = []
        else:
            tableData = tableData_raw or []
        
        # Load document
        doc = docx.Document(TEMPLATE_PATH)
        
        # 1. Format and populate Paragraphs
        # P0 (Title): Times New Roman, 18, Bold, Center aligned
        if len(doc.paragraphs) > 0:
            p0 = doc.paragraphs[0]
            p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p0.text = ""
            r0 = p0.add_run("LỊCH TUẦN CỬU – TIỂU TƯỜNG – ĐẠI TƯỜNG")
            r0.font.name = "Times New Roman"
            r0.font.size = Pt(18)
            r0.bold = True
            
        # P1 (Phẩm đạo/Name): Times New Roman, 16, R0 bold, R1 normal
        if len(doc.paragraphs) > 1:
            p1 = doc.paragraphs[1]
            p1.text = ""
            r0 = p1.add_run(phamDao)
            r0.font.name = "Times New Roman"
            r0.font.size = Pt(16)
            r0.bold = True
            
            if tuoiMat and tuoiMat > 0:
                details = f" : {hoTen} (Sinh năm {namSinhAl}, {tuoiTerm} {tuoiMat} tuổi)"
            else:
                details = f" : {hoTen}"
                
            r1 = p1.add_run(details)
            r1.font.name = "Times New Roman"
            r1.font.size = Pt(16)
                    
        # P2 (Lunar Death Date): Times New Roman, 16, R0 bold, R1 normal
        if len(doc.paragraphs) > 2:
            p2 = doc.paragraphs[2]
            p2.text = ""
            r0 = p2.add_run("Qui Vị/Qui Liễu")
            r0.font.name = "Times New Roman"
            r0.font.size = Pt(16)
            r0.bold = True
            r1 = p2.add_run(f": Ngày {ngayMatAl}, {gioMat} thời.")
            r1.font.name = "Times New Roman"
            r1.font.size = Pt(16)
                    
        # P3 (Solar Death Date): Times New Roman, 16, normal
        if len(doc.paragraphs) > 3:
            p3 = doc.paragraphs[3]
            p3.text = ""
            r0 = p3.add_run("    ")
            r0.font.name = "Times New Roman"
            r0.font.size = Pt(16)
            r1 = p3.add_run(f"Nhằm ngày: {ngayMatDl}")
            r1.font.name = "Times New Roman"
            r1.font.size = Pt(16)

        # P5 & P6: Format static text to Times New Roman, size 16, Center aligned
        for idx in [5, 6]:
            if idx < len(doc.paragraphs):
                p = doc.paragraphs[idx]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                text_runs = [r.text for r in p.runs]
                p.text = ""
                for t in text_runs:
                    r = p.add_run(t)
                    r.font.name = "Times New Roman"
                    r.font.size = Pt(16)
                    
        # 2. Fill in and format the table
        if doc.tables:
            table = doc.tables[0]
            
            # Format header row (Row 0)
            for cell in table.rows[0].cells:
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                text_runs = [r.text for r in p.runs]
                p.text = ""
                for t in text_runs:
                    r = p.add_run(t)
                    r.font.name = "Times New Roman"
                    r.font.size = Pt(16)
                    r.bold = True
            
            # Populate and format body rows (Row 1 to 11)
            for i, item in enumerate(tableData):
                if i + 1 < len(table.rows):
                    row = table.rows[i + 1]
                    
                    # Format column 0 in place
                    cell0 = row.cells[0]
                    for p in cell0.paragraphs:
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        text_runs = [r.text for r in p.runs]
                        p.text = ""
                        for t in text_runs:
                            r = p.add_run(t)
                            r.font.name = "Times New Roman"
                            r.font.size = Pt(16)
                    
                    # Populate columns 1, 2, 3
                    set_cell_text_and_style(row.cells[1], item['dayOfWeek'])
                    set_cell_text_and_style(row.cells[2], item['dateAl'])
                    set_cell_text_and_style(row.cells[3], item['dateDl'])
                    
        # Save document
        doc.save(OUTPUT_PATH)
        
        # Send file back
        return send_file(OUTPUT_PATH, as_attachment=True, download_name="KetQuaTuanCuu.docx")
        
    except Exception as e:
        print("Error generating document:", str(e))
        return jsonify({"error": str(e)}), 500

@app.route('/generate_causieu', methods=['POST'])
def generate_causieu():
    try:
        if request.is_json:
            data = request.json
        else:
            data = request.form
        val1 = data.get('val1', '')
        val2 = data.get('val2', '')
        val3 = data.get('val3', '')
        val4 = data.get('val4', '')
        val5 = data.get('val5', '')
        val6 = data.get('val6', '')
        val7 = data.get('val7', '')
        val8 = data.get('val8', '')
        val9 = data.get('val9', '')
        val10 = data.get('val10', '')
        cungXaType = data.get('cungXaType', 'xã')
        val11 = data.get('val11', '')
        huyenType = data.get('huyenType', 'huyện')
        val12 = data.get('val12', '')
        xaType = data.get('xaType', 'xã')
        val13 = data.get('val13', '')
        val14 = data.get('val14', '')
        val15 = data.get('val15', '')
        val16 = data.get('val16', '')
        filename = data.get('filename', 'Sớ cầu siêu')
        
        template_causieu_path = os.path.join(WORKSPACE, "CAU SIEU-2026.docx")
        doc = docx.Document(template_causieu_path)
        
        replacements = {
            "(1)": val1,
            "(2)": val2,
            "(3)": val3,
            "(4)": val4,
            "(5)": val5,
            "(6)": val6,
            "(7)": val7,
            "(8)": val8,
            "Nguyễn Thị Điếu": val8,
            "(9)": val9,
            "Bát thập Cửu": val9,
            "Bát Thập Cửu": val9,
            "Bát thập cửu": val9,
            "(10)": val10,
            "(11)": val11,
            "(12)": val12,
            "(13)": val13,
            "(14)": val14,
            "(15)": val15,
            "(16)": val16
        }
        
        for p in doc.paragraphs:
            for placeholder, value in replacements.items():
                p_text = p.text
                if placeholder in p_text:
                    start_char_idx = p_text.find(placeholder)
                    end_char_idx = start_char_idx + len(placeholder)
                    
                    current_len = 0
                    start_run_idx = -1
                    end_run_idx = -1
                    
                    for r_idx, r in enumerate(p.runs):
                        r_start = current_len
                        r_end = current_len + len(r.text)
                        current_len = r_end
                        
                        if r_start <= start_char_idx < r_end:
                            start_run_idx = r_idx
                        if r_start < end_char_idx <= r_end:
                            end_run_idx = r_idx
                            break
                    
                    if start_run_idx != -1 and end_run_idx != -1:
                        if start_run_idx == end_run_idx:
                            r = p.runs[start_run_idx]
                            r.text = r.text.replace(placeholder, value)
                            r.font.name = "Times New Roman"
                            r.font.color.rgb = docx.shared.RGBColor(0, 0, 0)
                        else:
                            run_start_len = sum(len(p.runs[j].text) for j in range(start_run_idx))
                            rel_start_idx = start_char_idx - run_start_len
                            
                            run_end_len = sum(len(p.runs[j].text) for j in range(end_run_idx))
                            rel_end_idx = end_char_idx - run_end_len
                            
                            r_start = p.runs[start_run_idx]
                            r_start.text = r_start.text[:rel_start_idx] + value
                            r_start.font.name = "Times New Roman"
                            r_start.font.color.rgb = docx.shared.RGBColor(0, 0, 0)
                            
                            for j in range(start_run_idx + 1, end_run_idx):
                                p.runs[j].text = ""
                                
                            r_end = p.runs[end_run_idx]
                            r_end.text = r_end.text[rel_end_idx:]

        # Specific replacements in P6 for cúng commune/ward type (after sáp nhập)
        if len(doc.paragraphs) > 6:
            p6 = doc.paragraphs[6]
            for r in p6.runs:
                if "xã" in r.text:
                    r.text = r.text.replace("xã", cungXaType)
                    r.font.name = "Times New Roman"
                    r.font.color.rgb = docx.shared.RGBColor(0, 0, 0)

        # Specific replacements in P16 for birthplace division types
        if len(doc.paragraphs) > 16:
            p16 = doc.paragraphs[16]
            for r in p16.runs:
                if "huyện" in r.text:
                    r.text = r.text.replace("huyện", huyenType)
                    r.font.name = "Times New Roman"
                    r.font.color.rgb = docx.shared.RGBColor(0, 0, 0)
                if "xã" in r.text:
                    r.text = r.text.replace("xã", xaType)
                    r.font.name = "Times New Roman"
                    r.font.color.rgb = docx.shared.RGBColor(0, 0, 0)

        temp_output_path = os.path.join(tempfile.gettempdir(), "Sớ cầu siêu tạm thời.docx")
        doc.save(temp_output_path)
        
        return send_file(temp_output_path, as_attachment=True, download_name=f"{filename}.docx")
        
    except Exception as e:
        print("Error generating Sớ cầu siêu:", str(e))
        return jsonify({"error": str(e)}), 500

@app.route('/generate_linhvi', methods=['POST'])
def generate_linhvi():
    try:
        if request.is_json:
            data = request.json
        else:
            data = request.form
            
        gioiTinh = data.get('gioiTinh', 'Nam')
        val1a = data.get('val1a', '')
        val1b = data.get('val1b', '')
        val2a = data.get('val2a', 'Bến')
        val2b = data.get('val2b', 'Tre')
        val3a = data.get('val3a', '')
        val3b = data.get('val3b', '')
        val4a = data.get('val4a', '')
        val4b = data.get('val4b', '')
        val5a = data.get('val5a', '')
        val5b = data.get('val5b', '')
        val6 = data.get('val6', '')
        val7 = data.get('val7', '')
        val8 = data.get('val8', '')
        val9 = data.get('val9', '')
        val10 = data.get('val10', '')
        wordThap = data.get('wordThap', 'Thập')
        val11a = data.get('val11a', '')
        val11b = data.get('val11b', '')
        val12 = data.get('val12', '')
        val13a = data.get('val13a', '')
        val13b = data.get('val13b', '')
        val14 = data.get('val14', '')
        val15a = data.get('val15a', '')
        val15b = data.get('val15b', '')
        
        # Decide template path based on gender
        if gioiTinh == 'Nữ':
            template_name = "MAU LINH VI NU.docx"
            filename_download = "Linh Vị Nữ"
        else:
            template_name = "MAU LINH VI NAM.docx"
            filename_download = "Linh Vị Nam"
            
        template_linhvi_path = os.path.join(WORKSPACE, template_name)
        doc = docx.Document(template_linhvi_path)
        
        replacements = {
            "(1a)": val1a,
            "(1b)": val1b,
            "(2a)": val2a,
            "(2b)": val2b,
            "(3a)": val3a,
            "(3b)": val3b,
            "(4a)": val4a,
            "(4b)": val4b,
            "(5a)": val5a,
            "(5b)": val5b,
            "(6)": val6,
            "(7)": val7,
            "(8)": val8,
            "(9)": val9,
            "(10)": val10,
            "THẬP": wordThap,
            "(11a)": val11a,
            "(11b)": val11b,
            "(12)": val12,
            "(13a)": val13a,
            "(13b)": val13b,
            "(14)": val14,
            "(15a)": val15a,
            "(15b)": val15b
        }
        
        placeholder_sizes = {
            "(1a)": 16, "(1b)": 16, "(2a)": 16, "(2b)": 16, "(3a)": 16, "(3b)": 16, "(4a)": 16, "(4b)": 16,
            "(5a)": 22, "(5b)": 22, "(6)": 22, "(7)": 22, "(8)": 22, "(9)": 22, "(10)": 22, "THẬP": 22,
            "(11a)": 15, "(11b)": 15, "(12)": 15, "(13a)": 15, "(13b)": 15, "(14)": 15, "(15a)": 15, "(15b)": 15
        }
        
        # Collect all paragraphs including inside tables
        all_paragraphs = list(doc.paragraphs)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    all_paragraphs.extend(cell.paragraphs)
                    
        for p in all_paragraphs:
            p_text = p.text
            for placeholder, value in replacements.items():
                if placeholder in p_text:
                    base_size = placeholder_sizes.get(placeholder, 16)
                    words = value.strip().split()
                    if len(words) >= 2:
                        if base_size == 22:
                            base_size = 16
                        elif base_size == 16:
                            base_size = 12
                        elif base_size == 15:
                            base_size = 11
                            
                    start_char_idx = p_text.find(placeholder)
                    end_char_idx = start_char_idx + len(placeholder)
                    
                    current_len = 0
                    start_run_idx = -1
                    end_run_idx = -1
                    
                    for r_idx, r in enumerate(p.runs):
                        r_start = current_len
                        r_end = current_len + len(r.text)
                        current_len = r_end
                        
                        if r_start <= start_char_idx < r_end:
                            start_run_idx = r_idx
                        if r_start < end_char_idx <= r_end:
                            end_run_idx = r_idx
                            break
                            
                    if start_run_idx != -1 and end_run_idx != -1:
                        if start_run_idx == end_run_idx:
                            r = p.runs[start_run_idx]
                            r.text = r.text.replace(placeholder, value.upper())
                            r.font.name = "Times New Roman"
                            r.font.size = Pt(base_size)
                            r.font.bold = False
                            r.font.color.rgb = docx.shared.RGBColor(0, 0, 0)
                        else:
                            run_start_len = sum(len(p.runs[j].text) for j in range(start_run_idx))
                            rel_start_idx = start_char_idx - run_start_len
                            
                            run_end_len = sum(len(p.runs[j].text) for j in range(end_run_idx))
                            rel_end_idx = end_char_idx - run_end_len
                            
                            r_start = p.runs[start_run_idx]
                            r_start.text = r_start.text[:rel_start_idx] + value.upper()
                            r_start.font.name = "Times New Roman"
                            r_start.font.size = Pt(base_size)
                            r_start.font.bold = False
                            r_start.font.color.rgb = docx.shared.RGBColor(0, 0, 0)
                            
                            for j in range(start_run_idx + 1, end_run_idx):
                                p.runs[j].text = ""
                                
                            r_end = p.runs[end_run_idx]
                            r_end.text = r_end.text[rel_end_idx:]
                            
                    p_text = p.text
                    
        temp_output_path = os.path.join(tempfile.gettempdir(), f"{filename_download}_tam_thoi.docx")
        doc.save(temp_output_path)
        
        return send_file(temp_output_path, as_attachment=True, download_name=f"{filename_download}.docx")
        
    except Exception as e:
        print("Error generating Linh Vị:", str(e))
        return jsonify({"error": str(e)}), 500

if __name__ == '__main__':
    app.run(host='127.0.0.1', port=5000, debug=True)
